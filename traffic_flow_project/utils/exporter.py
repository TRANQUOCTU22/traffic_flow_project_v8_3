from __future__ import annotations
from pathlib import Path
import csv
from typing import Dict
from docx import Document
from docx.shared import Pt
from models.graph_model import GraphModel, SolutionSummary


def build_report_text(graph_model:GraphModel, results:Dict[str,SolutionSummary], comparison_text:str)->str:
    # Tạo nội dung báo cáo tổng hợp.
    lines=[
        'BÁO CÁO KẾT QUẢ TỐI ƯU LUỒNG GIAO THÔNG',
        '='*70,
        f'Source: {graph_model.source}',
        f'Sink: {graph_model.sink}',
        f'Demand: {graph_model.demand}',
        '',
        '1. DANH SÁCH NODE VÀ EDGE'
    ]
    lines.append('Nodes: '+', '.join(graph_model.nodes))
    for e in graph_model.edges:
        lines.append(f'- {e.from_node}->{e.to_node} | capacity={e.capacity} | cost={e.cost}')

    if not results:
        lines += [
            '',
            '2. KẾT QUẢ THUẬT TOÁN',
            'Chưa chạy thuật toán nào nên chưa có bảng so sánh, nhận xét hoặc kết luận.'
        ]
        return '\n'.join(lines)

    lines.append('\n2. KẾT QUẢ TỪNG THUẬT TOÁN')
    for r in results.values():
        lines += [
            f'\n{r.algorithm}',
            f'Tổng luồng: {r.total_flow}',
            f'Tổng chi phí: {r.total_cost:.2f}',
            f'Thời gian trung bình: {r.runtime_seconds:.6f}s',
            f'Số lần chạy lấy trung bình: {r.extra.get("benchmark_runs",1)}',
            f'States: {r.extra.get("states_visited","-")}',
            f'Recursive calls: {r.extra.get("recursive_calls","-")}',
            f'Pruned: {r.extra.get("pruned_branches","-")}',
            'Path được chọn:'
        ]
        for p,f in r.path_flows:
            lines.append(f'  + {p.to_text()} | flow={f} | path_cost={p.cost}')
    lines += [
        '\n3. BẢNG SO SÁNH VÀ NHẬN XÉT',
        comparison_text,
        '\n4. KẾT LUẬN',
        'Basic Backtracking chậm do duyệt toàn bộ. Forward Checking giảm nhánh sai. Branch and Bound giảm mạnh số trạng thái. Heuristic giúp tìm nghiệm tốt sớm. Constraint Propagation thu hẹp miền giá trị. Full Optimization là phiên bản Backtracking hiệu quả nhất. Cuckoo Search nhanh hơn trên dữ liệu lớn nhưng không đảm bảo tối ưu tuyệt đối.'
    ]
    return '\n'.join(lines)


def export_txt(path,graph_model,results,comparison_text):
    # Ghi báo cáo ra file TXT.
    Path(path).write_text(build_report_text(graph_model,results,comparison_text),encoding='utf-8')


def export_csv(path,graph_model,results):
    # Ghi kết quả so sánh ra file CSV.
    with open(path,'w',newline='',encoding='utf-8-sig') as f:
        w=csv.writer(f)
        w.writerow(['Algorithm','Flow','Cost','Runtime avg','Runs','States','Calls','Pruned','Path count'])
        for r in results.values():
            w.writerow([
                r.algorithm,r.total_flow,r.total_cost,r.runtime_seconds,r.extra.get('benchmark_runs',1),
                r.extra.get('states_visited','-'),r.extra.get('recursive_calls','-'),r.extra.get('pruned_branches','-'),r.extra.get('path_count','-')
            ])


def export_docx(path,graph_model,results,comparison_text):
    # Ghi báo cáo ra file Word DOCX.
    doc=Document()
    doc.styles['Normal'].font.name='Arial'
    doc.styles['Normal'].font.size=Pt(10.5)
    doc.add_heading('Báo cáo kết quả tối ưu luồng giao thông',0)
    doc.add_paragraph(f'Source: {graph_model.source} - Sink: {graph_model.sink} - Demand: {graph_model.demand}')
    doc.add_heading('1. Dữ liệu đầu vào',1)
    table=doc.add_table(rows=1,cols=4); table.style='Table Grid'
    for c,t in zip(table.rows[0].cells,['From','To','Capacity','Cost']): c.text=t
    for e in graph_model.edges:
        row=table.add_row().cells; row[0].text=e.from_node; row[1].text=e.to_node; row[2].text=str(e.capacity); row[3].text=str(e.cost)

    if not results:
        doc.add_heading('2. Kết quả thuật toán',1)
        doc.add_paragraph('Chưa chạy thuật toán nào nên chưa có bảng so sánh, nhận xét hoặc kết luận.')
        doc.save(path)
        return

    doc.add_heading('2. Kết quả thuật toán',1)
    for r in results.values():
        doc.add_heading(r.algorithm,2)
        doc.add_paragraph(f'Flow={r.total_flow}, Cost={r.total_cost:.2f}, Runtime avg={r.runtime_seconds:.6f}s, Runs={r.extra.get("benchmark_runs",1)}')
        for p,f in r.path_flows:
            doc.add_paragraph(f'{p.to_text()} | flow={f} | cost={p.cost}',style='List Bullet')
    doc.add_heading('3. So sánh và nhận xét',1)
    doc.add_paragraph(comparison_text)
    doc.add_heading('4. Kết luận',1)
    doc.add_paragraph('Full Optimization là phiên bản Backtracking hiệu quả nhất trong nhóm exact search. Cuckoo Search phù hợp dữ liệu lớn hơn nhưng là metaheuristic nên không đảm bảo tối ưu tuyệt đối.')
    doc.save(path)
